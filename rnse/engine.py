"""
engine.py — Walk the document tree and perform placeholder substitution.

Handles:
  - Body paragraphs
  - Table cells (all rows × columns)
  - Document section headers and footers
"""

from __future__ import annotations

import logging
from typing import Generator, Tuple

from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

from rnse.formatter import format_value
from rnse.parser import merge_and_replace
from rnse.reporter import AuditReport

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Document tree walker
# ---------------------------------------------------------------------------

def iter_paragraph_groups(
    document: DocumentType,
) -> Generator[Tuple[Paragraph, str], None, None]:
    """
    Yield (paragraph, location_string) for every paragraph reachable in the
    document: body, tables, headers, and footers.

    Does NOT yield the paragraph if it has no runs (nothing to match).
    """
    # Body paragraphs
    for i, para in enumerate(document.paragraphs):
        yield para, f"paragraph:{i}"

    # Table cells
    for ti, table in enumerate(document.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, para in enumerate(cell.paragraphs):
                    yield para, f"table:{ti}:row:{ri}:col:{ci}:para:{pi}"

    # Headers and footers across all sections
    for si, section in enumerate(document.sections):
        for hf_name, hf_obj in [
            ("header", section.header),
            ("footer", section.footer),
            ("even_page_header", section.even_page_header),
            ("even_page_footer", section.even_page_footer),
            ("first_page_header", section.first_page_header),
            ("first_page_footer", section.first_page_footer),
        ]:
            if hf_obj is None:
                continue
            try:
                is_linked = hf_obj.is_linked_to_previous
            except Exception:
                is_linked = False
            if is_linked:
                continue
            try:
                paragraphs = hf_obj.paragraphs
            except Exception:
                continue
            for pi, para in enumerate(paragraphs):
                yield para, f"section:{si}:{hf_name}:para:{pi}"

            # Tables inside headers/footers
            try:
                tables = hf_obj.tables
            except Exception:
                tables = []
            for ti2, table in enumerate(tables):
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        for pi, para in enumerate(cell.paragraphs):
                            yield para, f"section:{si}:{hf_name}:table:{ti2}:row:{ri}:col:{ci}:para:{pi}"


# ---------------------------------------------------------------------------
# Main substitution entry point
# ---------------------------------------------------------------------------

def substitute_document(
    document: DocumentType,
    schedule: dict,
    reporter: AuditReport,
) -> None:
    """
    Walk the entire document tree, find all placeholders, and perform
    in-place substitution using *schedule*.

    *schedule* is a validated dict[asset_id][field] → float | None.
    *reporter* collects all substitutions and errors.

    After this call, *document* is mutated; save it with document.save().
    """
    log.debug("Starting document substitution")
    paragraph_count = 0
    match_count = 0

    for para, location in iter_paragraph_groups(document):
        runs = para.runs
        if not runs:
            continue

        # Quick check before entering merge_and_replace.
        combined = "".join(r.text or "" for r in runs)
        if "{{" not in combined:
            continue

        paragraph_count += 1
        before_subs = reporter.substitutions_ok
        before_issues = len(reporter.issues)

        merge_and_replace(
            runs=runs,
            schedule=schedule,
            formatter=format_value,
            reporter=reporter,
            location=location,
        )

        after_subs = reporter.substitutions_ok
        if after_subs > before_subs or len(reporter.issues) > before_issues:
            match_count += 1

    log.info(
        "Substitution complete: %d paragraphs with {{ processed, "
        "%d substitutions, %d errors, %d warnings",
        paragraph_count,
        reporter.substitutions_ok,
        reporter.error_count,
        reporter.warn_count,
    )

    # Warn about schedule assets that had no placeholders in the document.
    referenced_assets = {s.asset_id for s in reporter.substitutions}
    for asset_id in schedule:
        if asset_id not in referenced_assets:
            reporter.warn_unused_asset(asset_id)
