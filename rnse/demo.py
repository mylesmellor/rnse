"""
demo.py — Generate demo schedule.xlsx and demo_report.docx for immediate testing.

The demo report includes deliberate run-splitting to validate the parser's
run-merge strategy from the first test run.
"""

from __future__ import annotations

import logging
from pathlib import Path

import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Demo data
# ---------------------------------------------------------------------------

DEMO_ASSETS = [
    {
        "Asset_ID": "LON_001",
        "Asset_Name": "100 Bishopsgate, EC2",
        "MV": 2_500_000,
        "NIY": 0.0475,
        "TOPPED_UP_NIY": 0.0490,
        "RENT": 112_500,
        "ERV": 125_000,
        "AREA": 22_000,
        "CAPITAL_VALUE": 113.64,
    },
    {
        "Asset_ID": "LON_002",
        "Asset_Name": "45 Moorgate, EC2",
        "MV": 875_000,
        "NIY": 0.0550,
        "TOPPED_UP_NIY": 0.0565,
        "RENT": 45_000,
        "ERV": 50_000,
        "AREA": 8_500,
        "CAPITAL_VALUE": 102.94,
    },
    {
        "Asset_ID": "MCR_001",
        "Asset_Name": "1 Spinningfields, M3",
        "MV": 1_100_000,
        "NIY": 0.0600,
        "TOPPED_UP_NIY": 0.0615,
        "RENT": 63_000,
        "ERV": 70_000,
        "AREA": 14_000,
        "CAPITAL_VALUE": 78.57,
    },
]

FIELD_COLUMNS = ["MV", "NIY", "TOPPED_UP_NIY", "RENT", "ERV", "AREA", "CAPITAL_VALUE"]


# ---------------------------------------------------------------------------
# Schedule generator
# ---------------------------------------------------------------------------

def generate_demo_schedule(output_dir: Path, assets: list[dict] | None = None) -> Path:
    """Write demo/schedule.xlsx and return its path."""
    if assets is None:
        assets = DEMO_ASSETS

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Schedule"

    headers = ["Asset_ID", "Asset_Name"] + FIELD_COLUMNS
    ws.append(headers)

    for asset in assets:
        row = [asset.get(h) for h in headers]
        ws.append(row)

    # Basic column widths for readability.
    ws.column_dimensions["A"].width = 14
    ws.column_dimensions["B"].width = 28
    for col_letter in "CDEFGHI":
        ws.column_dimensions[col_letter].width = 16

    out_path = output_dir / "schedule.xlsx"
    wb.save(out_path)
    log.info("Demo schedule written to %s", out_path)
    return out_path


# ---------------------------------------------------------------------------
# Report generator
# ---------------------------------------------------------------------------

def _add_placeholder_split_runs(para, placeholder: str) -> None:
    """
    Add *placeholder* to *para* with deliberate run-splitting at ':' boundaries
    so the run-merge code is exercised.
    """
    # Split at ':' to create at least 3 runs per placeholder.
    parts = placeholder.split(":")
    for i, part in enumerate(parts):
        suffix = ":" if i < len(parts) - 1 else ""
        run = para.add_run(part + suffix)
        # Alternate bold on the separator runs to force Word to split runs.
        if i % 2 == 1:
            run.bold = True


def generate_demo_report(output_dir: Path, assets: list[dict] | None = None) -> Path:
    """Write demo/report.docx and return its path."""
    if assets is None:
        assets = DEMO_ASSETS

    doc = Document()

    # ---- Cover page --------------------------------------------------------
    title = doc.add_heading("Commercial Real Estate Valuation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Valuation Date: 31 January 2025")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---- Executive Summary -------------------------------------------------
    doc.add_heading("Executive Summary", level=1)

    exec_para = doc.add_paragraph(
        "This report summarises the valuation of a portfolio of three commercial "
        "properties located in London and Manchester. The aggregate portfolio market "
        "value is "
    )
    # LON_001 MV — deliberately split runs
    _add_placeholder_split_runs(exec_para, "{{MV:LON_001:£,0}}")
    exec_para.add_run(
        " for the Bishopsgate asset, "
    )
    _add_placeholder_split_runs(exec_para, "{{MV:LON_002:£,0}}")
    exec_para.add_run(
        " for the Moorgate asset, and "
    )
    _add_placeholder_split_runs(exec_para, "{{MV:MCR_001:£,0}}")
    exec_para.add_run(
        " for the Spinningfields asset."
    )

    doc.add_page_break()

    # ---- Individual asset sections -----------------------------------------
    for asset in assets:
        aid = asset["Asset_ID"]
        name = asset["Asset_Name"]

        doc.add_heading(name, level=2)

        # Narrative paragraph with run-split placeholders.
        para = doc.add_paragraph(
            f"The property at {name} has been assessed at a market value of "
        )
        _add_placeholder_split_runs(para, f"{{{{MV:{aid}:£,0}}}}")
        para.add_run(
            ", reflecting a net initial yield of "
        )
        _add_placeholder_split_runs(para, f"{{{{NIY:{aid}:0.00%}}}}")
        para.add_run(
            ". The passing rent is "
        )
        _add_placeholder_split_runs(para, f"{{{{RENT:{aid}:£,0}}}}")
        para.add_run(
            " per annum against an estimated rental value of "
        )
        _add_placeholder_split_runs(para, f"{{{{ERV:{aid}:£,0}}}}")
        para.add_run(".")

        # Summary table for this asset.
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        def _set_cell(r: int, c: int, text: str, bold: bool = False) -> None:
            cell = table.cell(r, c)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(text)
            run.bold = bold

        _set_cell(0, 0, "Field", bold=True)
        _set_cell(0, 1, "Value", bold=True)
        _set_cell(1, 0, "Floor Area")
        _set_cell(2, 0, "Capital Value")
        _set_cell(3, 0, "Topped-Up NIY")

        # Placeholders in table cells — these have their own run context.
        for row_idx, (field, fmt) in enumerate([
            ("AREA", f"#,##0 sq ft"),
            ("CAPITAL_VALUE", "psf"),
            ("TOPPED_UP_NIY", "0.00%"),
        ], start=1):
            cell = table.cell(row_idx, 1)
            cell.paragraphs[0].clear()
            _add_placeholder_split_runs(
                cell.paragraphs[0], f"{{{{{field}:{aid}:{fmt}}}}}"
            )

        doc.add_paragraph()  # Spacer

    doc.add_page_break()

    # ---- Portfolio Summary Table -------------------------------------------
    doc.add_heading("Portfolio Summary", level=1)

    all_fields = [("MV", "£,0"), ("NIY", "0.00%"), ("RENT", "£,0"), ("ERV", "£,0"),
                  ("AREA", "#,##0 sq ft"), ("CAPITAL_VALUE", "psf")]

    n_cols = len(all_fields) + 1
    port_table = doc.add_table(rows=len(assets) + 1, cols=n_cols)
    port_table.style = "Table Grid"

    # Header row.
    port_table.cell(0, 0).paragraphs[0].add_run("Asset").bold = True
    for ci, (fn, _) in enumerate(all_fields, start=1):
        port_table.cell(0, ci).paragraphs[0].add_run(fn).bold = True

    # Data rows — normal (non-split) placeholders to test the simple path.
    for ri, asset in enumerate(assets, start=1):
        aid = asset["Asset_ID"]
        port_table.cell(ri, 0).paragraphs[0].add_run(asset["Asset_Name"])
        for ci, (fn, fmt) in enumerate(all_fields, start=1):
            cell = port_table.cell(ri, ci)
            cell.paragraphs[0].add_run(f"{{{{{fn}:{aid}:{fmt}}}}}")

    # ---- Footer with placeholder -------------------------------------------
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.clear()
    footer_para.add_run("Flagship asset market value: ")
    # Use a normal (non-split) run in the footer.
    footer_para.add_run("{{MV:LON_001:£m}}")
    footer_para.add_run("  |  Confidential")

    out_path = output_dir / "report.docx"
    doc.save(str(out_path))
    log.info("Demo report written to %s", out_path)
    return out_path
