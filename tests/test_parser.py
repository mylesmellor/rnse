"""
test_parser.py — Tests for placeholder regex and run-merging logic.
"""

from __future__ import annotations

import pytest
from docx import Document

from rnse.parser import (
    PLACEHOLDER_RE,
    PlaceholderToken,
    build_run_spans,
    apply_replacement,
    parse_placeholders,
    has_malformed_placeholder,
)


# ---------------------------------------------------------------------------
# Regex / parse_placeholders
# ---------------------------------------------------------------------------

class TestParsePlaceholders:
    def test_simple_match(self):
        tokens = parse_placeholders("{{MV:ASSET_001:£,0}}")
        assert len(tokens) == 1
        t = tokens[0]
        assert t.field == "MV"
        assert t.asset_id == "ASSET_001"
        assert t.format_spec == "£,0"
        assert t.raw == "{{MV:ASSET_001:£,0}}"
        assert t.start == 0
        assert t.end == len("{{MV:ASSET_001:£,0}}")

    def test_placeholder_at_start(self):
        tokens = parse_placeholders("{{NIY:ASSET_001:0.00%}} is the yield")
        assert len(tokens) == 1
        assert tokens[0].start == 0

    def test_placeholder_at_end(self):
        tokens = parse_placeholders("The yield is {{NIY:ASSET_001:0.00%}}")
        assert len(tokens) == 1
        assert tokens[0].end == len("The yield is {{NIY:ASSET_001:0.00%}}")

    def test_multiple_placeholders(self):
        text = "Value: {{MV:A001:£,0}} Yield: {{NIY:A001:0.00%}} Area: {{AREA:A001:#,##0 sq ft}}"
        tokens = parse_placeholders(text)
        assert len(tokens) == 3
        assert tokens[0].field == "MV"
        assert tokens[1].field == "NIY"
        assert tokens[2].field == "AREA"

    def test_no_false_positive_single_braces(self):
        tokens = parse_placeholders("{single braces} are not tokens")
        assert tokens == []

    def test_no_false_positive_partial_open(self):
        tokens = parse_placeholders("{{partial open without close")
        assert tokens == []

    def test_no_false_positive_empty_field(self):
        # Fields must start with uppercase letter.
        tokens = parse_placeholders("{{:ASSET_001:£,0}}")
        assert tokens == []

    def test_underscored_field(self):
        tokens = parse_placeholders("{{TOPPED_UP_NIY:ASSET_001:0.00%}}")
        assert len(tokens) == 1
        assert tokens[0].field == "TOPPED_UP_NIY"

    def test_asset_id_with_digits(self):
        tokens = parse_placeholders("{{MV:LON001:£m}}")
        assert len(tokens) == 1
        assert tokens[0].asset_id == "LON001"

    def test_asset_id_with_underscore(self):
        tokens = parse_placeholders("{{MV:LON_001:£m}}")
        assert len(tokens) == 1
        assert tokens[0].asset_id == "LON_001"

    def test_format_spec_with_suffix(self):
        tokens = parse_placeholders("{{AREA:ASSET_001:#,##0 sq ft}}")
        assert len(tokens) == 1
        assert tokens[0].format_spec == "#,##0 sq ft"

    def test_psf_format_spec(self):
        tokens = parse_placeholders("{{CAPITAL_VALUE:ASSET_001:psf}}")
        assert len(tokens) == 1
        assert tokens[0].format_spec == "psf"

    def test_position_tracking(self):
        prefix = "Value is "
        placeholder = "{{MV:ASSET_001:£,0}}"
        text = prefix + placeholder
        tokens = parse_placeholders(text)
        assert tokens[0].start == len(prefix)
        assert tokens[0].end == len(text)

    def test_adjacent_placeholders(self):
        text = "{{MV:A:£,0}}{{NIY:A:0.00%}}"
        tokens = parse_placeholders(text)
        assert len(tokens) == 2
        assert tokens[1].start == tokens[0].end


# ---------------------------------------------------------------------------
# has_malformed_placeholder
# ---------------------------------------------------------------------------

class TestMalformedDetection:
    def test_no_malformed_when_valid(self):
        text = "Value: {{MV:ASSET_001:£,0}}"
        tokens = parse_placeholders(text)
        assert not has_malformed_placeholder(text, tokens)

    def test_detects_lone_double_brace(self):
        text = "Something {{ broken here"
        tokens = parse_placeholders(text)
        assert has_malformed_placeholder(text, tokens)

    def test_detects_partial_placeholder(self):
        text = "{{MV:ASSET_001  missing close"
        tokens = parse_placeholders(text)
        assert has_malformed_placeholder(text, tokens)

    def test_no_double_brace_no_warning(self):
        text = "No placeholders here at all"
        tokens = parse_placeholders(text)
        assert not has_malformed_placeholder(text, tokens)


# ---------------------------------------------------------------------------
# build_run_spans + apply_replacement
# ---------------------------------------------------------------------------

class TestRunMerging:
    def _make_runs(self, texts: list[str]):
        """Create a minimal paragraph with the given run texts."""
        doc = Document()
        para = doc.add_paragraph()
        for element in list(para._element):
            para._element.remove(element)
        for t in texts:
            para.add_run(t)
        return para.runs

    def test_build_run_spans_single_run(self):
        runs = self._make_runs(["hello world"])
        combined, spans = build_run_spans(runs)
        assert combined == "hello world"
        assert len(spans) == 1
        assert spans[0].char_start == 0
        assert spans[0].char_end == 11

    def test_build_run_spans_multiple_runs(self):
        runs = self._make_runs(["hello", " ", "world"])
        combined, spans = build_run_spans(runs)
        assert combined == "hello world"
        assert spans[0].char_end == 5
        assert spans[1].char_start == 5
        assert spans[1].char_end == 6
        assert spans[2].char_start == 6
        assert spans[2].char_end == 11

    def test_apply_replacement_single_run(self):
        runs = self._make_runs(["{{MV:ASSET_001:£,0}}"])
        combined, spans = build_run_spans(runs)
        tokens = parse_placeholders(combined)
        apply_replacement(runs, tokens[0], "£1,250,000", spans)
        assert runs[0].text == "£1,250,000"

    def test_apply_replacement_three_runs(self):
        # Placeholder split across 3 runs: "{{MV" | ":ASSET_001" | ":£,0}}"
        runs = self._make_runs(["{{MV", ":ASSET_001", ":£,0}}"])
        combined, spans = build_run_spans(runs)
        tokens = parse_placeholders(combined)
        assert len(tokens) == 1
        apply_replacement(runs, tokens[0], "£1,250,000", spans)
        assert runs[0].text == "£1,250,000"
        assert runs[1].text == ""
        assert runs[2].text == ""

    def test_apply_replacement_preserves_prefix_suffix(self):
        # "Value: {{MV:A:£,0}} approx" in one run.
        runs = self._make_runs(["Value: {{MV:A:£,0}} approx"])
        combined, spans = build_run_spans(runs)
        tokens = parse_placeholders(combined)
        apply_replacement(runs, tokens[0], "£100", spans)
        assert runs[0].text == "Value: £100 approx"

    def test_apply_replacement_split_with_prefix(self):
        # "Value: {{MV" in run 0, ":A:£,0}} end" in run 1.
        runs = self._make_runs(["Value: {{MV", ":A:£,0}} end"])
        combined, spans = build_run_spans(runs)
        tokens = parse_placeholders(combined)
        apply_replacement(runs, tokens[0], "£100", spans)
        assert runs[0].text == "Value: £100"
        assert runs[1].text == " end"

    def test_apply_replacement_multiple_in_one_paragraph(self):
        # Two placeholders in a single run.
        runs = self._make_runs(["{{MV:A:£,0}} and {{NIY:A:0.00%}}"])
        combined, spans = build_run_spans(runs)
        tokens = parse_placeholders(combined)
        assert len(tokens) == 2
        # Process in reverse order.
        for token in reversed(tokens):
            replacement = "£100" if token.field == "MV" else "5.25%"
            apply_replacement(runs, token, replacement, spans)
        # After both replacements, combined text should reflect both.
        result = "".join(r.text for r in runs)
        assert "£100" in result
        assert "5.25%" in result
