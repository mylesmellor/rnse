"""
conftest.py — Shared fixtures for the RNSE test suite.

Provides:
  make_docx(paragraphs)     — build a minimal in-memory docx
  make_schedule(rows)       — build a simple schedule dict
  simple_schedule           — pre-built schedule fixture
"""

from __future__ import annotations

import io
from typing import Sequence

import pytest
from docx import Document
from docx.document import Document as DocumentType


# ---------------------------------------------------------------------------
# Schedule builder
# ---------------------------------------------------------------------------

def make_schedule(
    rows: list[dict],
    fields: list[str] | None = None,
) -> dict:
    """
    Build a validated schedule dict from a list of dicts.

    Each dict should contain 'asset_id' plus any field keys.
    Returns a dict[asset_id][field] → float | None.
    """
    if fields is None and rows:
        # Infer from first row, excluding 'asset_id'.
        fields = [k for k in rows[0] if k.upper() != "ASSET_ID"]

    schedule: dict = {}
    for row in rows:
        aid = row["asset_id"].upper()
        schedule[aid] = {
            k.upper(): float(v) if v is not None else None
            for k, v in row.items()
            if k.upper() != "ASSET_ID"
        }
    return schedule


# ---------------------------------------------------------------------------
# Docx builder helpers
# ---------------------------------------------------------------------------

def _make_split_runs(para, text: str, split_at: str = ":") -> None:
    """Add *text* to *para*, splitting runs at each *split_at* character."""
    parts = text.split(split_at)
    for i, part in enumerate(parts):
        suffix = split_at if i < len(parts) - 1 else ""
        run = para.add_run(part + suffix)
        if i % 2 == 1:
            run.bold = True


def make_docx(paragraphs: Sequence[str | list[str]]) -> DocumentType:
    """
    Build a minimal in-memory Document.

    *paragraphs* is a list of items:
      - str: added as a single run
      - list[str]: each element added as a separate run (simulates run-splitting)
    """
    doc = Document()
    # Remove the default empty paragraph.
    for element in list(doc.element.body):
        doc.element.body.remove(element)

    for item in paragraphs:
        para = doc.add_paragraph()
        if isinstance(item, str):
            para.add_run(item)
        else:
            for part in item:
                para.add_run(part)
    return doc


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def simple_schedule() -> dict:
    return make_schedule([
        {
            "asset_id": "ASSET_001",
            "MV": 1_250_000.0,
            "NIY": 0.0525,
            "RENT": 75_000.0,
            "ERV": 80_000.0,
            "AREA": 10_500.0,
            "CAPITAL_VALUE": 119.05,
            "TOPPED_UP_NIY": 0.0550,
        },
        {
            "asset_id": "ASSET_002",
            "MV": 4_750_000.0,
            "NIY": 0.0480,
            "RENT": 220_000.0,
            "ERV": 235_000.0,
            "AREA": 42_000.0,
            "CAPITAL_VALUE": 113.10,
            "TOPPED_UP_NIY": 0.0495,
        },
    ])
