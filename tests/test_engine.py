"""
test_engine.py — Integration tests for full document substitution.
"""

from __future__ import annotations

import pytest
from docx import Document

from rnse.engine import substitute_document
from rnse.reporter import AuditReport
from tests.conftest import make_docx, make_schedule


def _run_engine(doc, schedule):
    """Helper: run engine and return reporter."""
    reporter = AuditReport()
    substitute_document(doc, schedule, reporter)
    return reporter


def _text(doc) -> str:
    """Concatenate all paragraph text in the document body."""
    return "\n".join(p.text for p in doc.paragraphs)


def _all_text(doc) -> str:
    """Concatenate all paragraph text including tables."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    parts.append(para.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Basic substitution
# ---------------------------------------------------------------------------

class TestSingleParagraphSinglePlaceholder:
    def test_substitution_performed(self, simple_schedule):
        doc = make_docx(["Market value: {{MV:ASSET_001:£,0}}"])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 1
        assert "£1,250,000" in _text(doc)
        assert "{{" not in _text(doc)

    def test_substitution_logged(self, simple_schedule):
        doc = make_docx(["{{MV:ASSET_001:£,0}}"])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions[0].placeholder == "{{MV:ASSET_001:£,0}}"
        assert reporter.substitutions[0].formatted_value == "£1,250,000"
        assert reporter.substitutions[0].raw_value == 1_250_000.0

    def test_yields_formatted(self, simple_schedule):
        doc = make_docx(["Yield: {{NIY:ASSET_001:0.00%}}"])
        _run_engine(doc, simple_schedule)
        assert "5.25%" in _text(doc)

    def test_area_with_suffix(self, simple_schedule):
        doc = make_docx(["Area: {{AREA:ASSET_001:#,##0 sq ft}}"])
        _run_engine(doc, simple_schedule)
        assert "10,500 sq ft" in _text(doc)

    def test_millions_format(self, simple_schedule):
        doc = make_docx(["{{MV:ASSET_001:£m}}"])
        _run_engine(doc, simple_schedule)
        assert "£1.3m" in _text(doc)


class TestMultiplePlaceholdersOneParagraph:
    def test_all_replaced(self, simple_schedule):
        doc = make_docx(["MV={{MV:ASSET_001:£,0}} NIY={{NIY:ASSET_001:0.00%}}"])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 2
        t = _text(doc)
        assert "£1,250,000" in t
        assert "5.25%" in t

    def test_ten_consecutive_placeholders(self, simple_schedule):
        tokens = " ".join([f"{{{{MV:ASSET_001:£,0}}}}" for _ in range(10)])
        doc = make_docx([tokens])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 10
        assert "{{" not in _text(doc)


# ---------------------------------------------------------------------------
# Run-split tests
# ---------------------------------------------------------------------------

class TestRunSplitting:
    def test_placeholder_split_across_two_runs(self, simple_schedule):
        # "{{MV:ASSET" in run 0, "_001:£,0}}" in run 1
        doc = make_docx([["{{MV:ASSET", "_001:£,0}}"]])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 1
        assert "£1,250,000" in _text(doc)

    def test_placeholder_split_across_three_runs(self, simple_schedule):
        doc = make_docx([["{{MV", ":ASSET_001", ":£,0}}"]])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 1
        assert "£1,250,000" in _text(doc)

    def test_separator_colon_in_separate_run(self, simple_schedule):
        # Colon between FIELD and ASSET_ID in a separate run.
        doc = make_docx([["{{MV:", "ASSET_001", ":£,0}}"]])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 1

    def test_split_preserves_surrounding_text(self, simple_schedule):
        doc = make_docx([["Value is ", "{{MV:ASSET", "_001:£,0}}", " end"]])
        _run_engine(doc, simple_schedule)
        t = _text(doc)
        assert t.startswith("Value is ")
        assert t.endswith(" end")
        assert "£1,250,000" in t


# ---------------------------------------------------------------------------
# Table cells
# ---------------------------------------------------------------------------

class TestTableCells:
    def _make_doc_with_table(self, schedule):
        doc = Document()
        # Do NOT clear the body: removing all elements strips the sectPr,
        # which breaks doc.sections. Just add the table alongside the default paragraph.
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).paragraphs[0].add_run("MV")
        table.cell(0, 1).paragraphs[0].add_run("{{MV:ASSET_001:£,0}}")
        table.cell(1, 0).paragraphs[0].add_run("NIY")
        table.cell(1, 1).paragraphs[0].add_run("{{NIY:ASSET_001:0.00%}}")
        return doc

    def test_table_cell_substitution(self, simple_schedule):
        doc = self._make_doc_with_table(simple_schedule)
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 2
        text = _all_text(doc)
        assert "£1,250,000" in text
        assert "5.25%" in text

    def test_table_cell_location_in_audit(self, simple_schedule):
        doc = self._make_doc_with_table(simple_schedule)
        reporter = _run_engine(doc, simple_schedule)
        locations = [s.location for s in reporter.substitutions]
        assert any("table" in loc for loc in locations)


# ---------------------------------------------------------------------------
# Headers and footers
# ---------------------------------------------------------------------------

class TestHeaderFooter:
    def _make_doc_with_footer(self, placeholder: str):
        doc = Document()
        # Do NOT clear the body: it strips the sectPr and breaks doc.sections.
        doc.add_paragraph("Body text only")
        footer = doc.sections[0].footer
        footer.paragraphs[0].add_run(placeholder)
        return doc

    def test_footer_substitution(self, simple_schedule):
        doc = self._make_doc_with_footer("{{MV:ASSET_001:£m}}")
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 1
        footer_text = doc.sections[0].footer.paragraphs[0].text
        assert "£1.3m" in footer_text

    def test_footer_location_in_audit(self, simple_schedule):
        doc = self._make_doc_with_footer("{{MV:ASSET_001:£m}}")
        reporter = _run_engine(doc, simple_schedule)
        assert any("footer" in s.location for s in reporter.substitutions)


# ---------------------------------------------------------------------------
# Error handling
# ---------------------------------------------------------------------------

class TestErrorHandling:
    def test_unknown_asset_id_error_logged(self, simple_schedule):
        doc = make_docx(["{{MV:ASSET_999:£,0}}"])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.error_count == 1
        assert reporter.issues[0].code == "ERROR_UNKNOWN_ASSET_ID"

    def test_unknown_asset_id_placeholder_left_unchanged(self, simple_schedule):
        doc = make_docx(["{{MV:ASSET_999:£,0}}"])
        _run_engine(doc, simple_schedule)
        assert "{{MV:ASSET_999:£,0}}" in _text(doc)

    def test_unknown_field_error_logged(self, simple_schedule):
        doc = make_docx(["{{NONEXISTENT_FIELD:ASSET_001:£,0}}"])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.error_count == 1
        assert reporter.issues[0].code == "ERROR_UNKNOWN_FIELD"

    def test_missing_value_error_logged(self):
        schedule = make_schedule([{"asset_id": "ASSET_001", "MV": None}])
        doc = make_docx(["{{MV:ASSET_001:£,0}}"])
        reporter = _run_engine(doc, schedule)
        assert reporter.error_count == 1
        assert reporter.issues[0].code == "ERROR_MISSING_VALUE"

    def test_bad_format_spec_error_logged(self, simple_schedule):
        doc = make_docx(["{{MV:ASSET_001:bad_spec}}"])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.error_count == 1
        assert reporter.issues[0].code == "ERROR_UNKNOWN_FORMAT_SPEC"

    def test_continues_after_error(self, simple_schedule):
        doc = make_docx(["{{MV:ASSET_999:£,0}} then {{NIY:ASSET_001:0.00%}}"])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.error_count == 1
        assert reporter.substitutions_ok == 1

    def test_unused_asset_warning(self, simple_schedule):
        # Document only references ASSET_001 but schedule has ASSET_002 too.
        doc = make_docx(["{{MV:ASSET_001:£,0}}"])
        reporter = _run_engine(doc, simple_schedule)
        warn_codes = [i.code for i in reporter.issues if i.severity == "WARN"]
        assert "WARN_UNUSED_ASSET" in warn_codes


class TestNoPlaceholders:
    def test_clean_document_no_substitutions(self, simple_schedule):
        doc = make_docx(["No placeholders here.", "Just plain text."])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 0
        assert reporter.error_count == 0

    def test_clean_document_unused_asset_warnings(self, simple_schedule):
        doc = make_docx(["No placeholders here."])
        reporter = _run_engine(doc, simple_schedule)
        warn_codes = [i.code for i in reporter.issues]
        assert "WARN_UNUSED_ASSET" in warn_codes


# ---------------------------------------------------------------------------
# Repeated placeholders
# ---------------------------------------------------------------------------

class TestRepetition:
    def test_same_placeholder_multiple_paragraphs(self, simple_schedule):
        doc = make_docx([
            "First: {{MV:ASSET_001:£,0}}",
            "Second: {{MV:ASSET_001:£,0}}",
            "Third: {{MV:ASSET_001:£,0}}",
        ])
        reporter = _run_engine(doc, simple_schedule)
        assert reporter.substitutions_ok == 3
        t = _text(doc)
        assert t.count("£1,250,000") == 3
